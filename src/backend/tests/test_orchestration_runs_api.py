from io import BytesIO

from docx import Document

from app.api.routes import orchestration as orchestration_route


class _FakeTask:
    id = "celery-task-1"


def _build_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Senior Frontend Engineer")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_create_orchestration_run_returns_run_metadata(client, monkeypatch) -> None:
    events: list[dict[str, object]] = []
    run_state: dict[str, object] = {}

    def fake_create_run(*, run_id: str, candidate_id: str, filename: str) -> None:
        run_state.update(
            {
                "run_id": run_id,
                "candidate_id": candidate_id,
                "filename": filename,
                "status": "queued",
            }
        )

    def fake_set_status(run_id: str, *, status: str, result=None, error: str = "", task_id: str = "") -> None:
        run_state.update({"run_id": run_id, "status": status, "task_id": task_id, "error": error, "result": result})

    def fake_publish_event(run_id: str, event: dict[str, object]) -> dict[str, object]:
        payload = {"sequence": len(events) + 1, **event}
        events.append(payload)
        return payload

    monkeypatch.setattr(orchestration_route.orchestration_progress_service, "create_run", fake_create_run)
    monkeypatch.setattr(orchestration_route.orchestration_progress_service, "set_status", fake_set_status)
    monkeypatch.setattr(orchestration_route.orchestration_progress_service, "publish_event", fake_publish_event)
    monkeypatch.setattr(orchestration_route.task_queue, "enqueue_candidate_analysis", lambda **kwargs: _FakeTask())

    response = client.post(
        "/api/orchestration/runs",
        files={
            "file": (
                "resume.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"job_title": "Senior Frontend Engineer"},
    )

    assert response.status_code == 200
    body = response.json()
    assert body["task_id"] == "celery-task-1"
    assert body["run_id"].startswith("analysis-")
    assert events[0]["type"] == "run_queued"


def test_orchestration_events_endpoint_replays_existing_backlog(client, monkeypatch) -> None:
    backlog = [
        {
            "sequence": 1,
            "run_id": "analysis-1",
            "type": "run_started",
            "label": "candidate_analysis",
            "stage": "candidate_analysis",
            "timestamp": "2026-07-06T00:00:00Z",
        },
        {
            "sequence": 2,
            "run_id": "analysis-1",
            "type": "run_completed",
            "label": "candidate_analysis",
            "stage": "candidate_analysis",
            "timestamp": "2026-07-06T00:00:01Z",
        },
    ]

    monkeypatch.setattr(
        orchestration_route.orchestration_progress_service,
        "get_run",
        lambda run_id: {"run_id": run_id, "status": "completed", "result": {"ok": True}, "error": ""},
    )
    monkeypatch.setattr(
        orchestration_route.orchestration_progress_service,
        "get_events",
        lambda run_id: backlog,
    )

    response = client.get("/api/orchestration/runs/analysis-1/events")

    assert response.status_code == 200
    assert '"type": "run_started"' in response.text
    assert '"type": "run_completed"' in response.text
