from io import BytesIO

from docx import Document

from app.api.deps import task_queue


def test_enqueue_cv_extraction(client) -> None:
    class DummyTask:
        id = "task-123"
        state = "PENDING"

    original_enqueue = task_queue.enqueue_cv_extraction
    task_queue.enqueue_cv_extraction = lambda **kwargs: DummyTask()
    try:
        document = Document()
        document.add_paragraph("Jane Doe")
        document.add_paragraph("Python Engineer")
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        response = client.post(
            "/api/cv/extract",
            files={
                "file": (
                    "resume.docx",
                    buffer.read(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert response.status_code == 200
        payload = response.json()
        assert payload["task_id"] == "task-123"
        assert payload["candidate_id"]
    finally:
        task_queue.enqueue_cv_extraction = original_enqueue
