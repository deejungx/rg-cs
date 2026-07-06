from io import BytesIO

from docx import Document

from app.core.config import settings


def test_candidate_analysis_api_returns_trace_and_mock_mode(client) -> None:
    settings.openai_api_key = ""
    settings.model_provider = "auto"

    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Senior Frontend Engineer")
    document.add_paragraph("Email: jane@example.com")
    document.add_paragraph("Skills: React, TypeScript, Tailwind, Accessibility")
    document.add_paragraph("Experience")
    document.add_paragraph("Built React interfaces and design systems.")

    content = BytesIO()
    document.save(content)
    content.seek(0)

    response = client.post(
        "/api/orchestration/analyze",
        files={
            "file": (
                "resume.docx",
                content.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "job_title": "Senior Frontend Engineer",
            "job_description": "Build React and TypeScript interfaces for a hiring platform. Hybrid full-time role.",
            "company_name": "ABC Company",
            "location": "Kathmandu",
            "skills_csv": "React, TypeScript, Tailwind",
        },
    )

    assert response.status_code == 200
    body = response.json()
    assert body["model_provider"] == "mock"
    assert body["model_mode"] == "mock"
    assert body["trace"][0]["step"] == "cv_extraction"
    assert body["trace"][0]["input_contract"]["name"] == "cv_extraction.input"
    assert body["extraction"]["candidate_id"] == body["candidate_id"]
    assert body["vacancy"]["title"] == "Senior Frontend Engineer"
