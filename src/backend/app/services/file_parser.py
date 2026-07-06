from pathlib import Path

import docx
from pypdf import PdfReader

from src.shared.schemas import DocumentParseResult


class FileParserService:
    def parse_path(self, source_path: Path) -> DocumentParseResult:
        suffix = source_path.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(source_path)
        if suffix == ".docx":
            return self._parse_docx(source_path)
        if suffix in {".txt", ".md"}:
            return self._parse_text(source_path)
        raise ValueError(f"Unsupported file type: {suffix}")

    def _parse_pdf(self, source_path: Path) -> DocumentParseResult:
        reader = PdfReader(str(source_path))
        parts = [(page.extract_text() or "").strip() for page in reader.pages]
        text = "\n\n".join(part for part in parts if part)
        result = DocumentParseResult(
            text=text,
            parser="pypdf",
            file_type="pdf",
            is_image_based=not bool(text.strip()),
        )
        result.set_source_text(text)
        return result

    def _parse_docx(self, source_path: Path) -> DocumentParseResult:
        document = docx.Document(str(source_path))
        lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    lines.append(" | ".join(values))
        text = "\n".join(lines)
        result = DocumentParseResult(
            text="\n".join(lines),
            parser="python-docx",
            file_type="docx",
            is_image_based=False,
        )
        result.set_source_text(text)
        return result

    def _parse_text(self, source_path: Path) -> DocumentParseResult:
        text = source_path.read_text(encoding="utf-8")
        suffix = source_path.suffix.lower().lstrip(".")
        result = DocumentParseResult(
            text=text,
            parser="text",
            file_type=suffix,
            is_image_based=False,
        )
        result.set_source_text(text)
        return result


file_parser_service = FileParserService()
