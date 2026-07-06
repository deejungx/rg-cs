import json
from io import BytesIO
from pathlib import Path

from docx import Document

from app.core.config import settings
from app.services.upload_service import upload_service
from app.services.workspace_service import workspace_service
from src.ai.pipelines.cv_extraction_flow import run_cv_extraction_flow


def test_cv_flow_redacts_parsed_text_in_result_and_trace(tmp_path) -> None:
    settings.uploads_dir = tmp_path / "uploads"
    settings.workspace_dir = tmp_path / "workspace"
    upload_service.uploads_dir = settings.uploads_dir
    workspace_service.uploads_dir = settings.uploads_dir
    workspace_service.workspace_dir = settings.workspace_dir
    workspace_service.ensure_runtime_dirs()

    candidate_id = "candidate-pii"
    upload_dir = settings.uploads_dir / "resumes" / candidate_id
    upload_dir.mkdir(parents=True, exist_ok=True)
    source_path = upload_dir / "resume.docx"

    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Email: jane.doe@example.com")
    document.add_paragraph("Phone: +1 555 123 4567")
    document.add_paragraph("Portfolio: https://example.com/jane-doe")
    document.add_paragraph("Experience with FastAPI and Redis")
    document.save(BytesIO())
    document.save(source_path)

    result = run_cv_extraction_flow(
        candidate_id=candidate_id,
        source_path=str(source_path),
        filename="resume.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    parsed_text = result["parsed"]["text"]
    assert "jane.doe@example.com" not in parsed_text
    assert "+1 555 123 4567" not in parsed_text
    assert "https://example.com/jane-doe" not in parsed_text
    assert result["parsed"]["redaction_applied"] is True

    trace_path = settings.workspace_dir / Path(result["artifacts"]["trace_path"])
    trace_payload = json.loads(trace_path.read_text(encoding="utf-8"))
    trace_text = trace_payload["parser"]["text"]

    assert "jane.doe@example.com" not in trace_text
    assert "+1 555 123 4567" not in trace_text
    assert "https://example.com/jane-doe" not in trace_text
