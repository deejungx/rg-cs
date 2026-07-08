from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from app.core.config import settings
from app.services.upload_service import upload_service
from app.services.workspace_service import workspace_service
from src.ai.pipelines.cv_extraction_flow import InvalidResumeError, run_cv_extraction_flow


def test_cv_flow_creates_workspace_artifacts(tmp_path) -> None:
    settings.uploads_dir = tmp_path / "uploads"
    settings.workspace_dir = tmp_path / "workspace"
    upload_service.uploads_dir = settings.uploads_dir
    workspace_service.uploads_dir = settings.uploads_dir
    workspace_service.workspace_dir = settings.workspace_dir
    workspace_service.ensure_runtime_dirs()

    candidate_id = "candidate-test"
    upload_dir = settings.uploads_dir / "resumes" / candidate_id
    upload_dir.mkdir(parents=True, exist_ok=True)
    source_path = upload_dir / "resume.docx"

    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("jane@example.com")
    document.add_paragraph("Senior Python Engineer")
    document.add_paragraph("Experience with FastAPI, Redis, Celery, and Qdrant")
    document.add_paragraph("Skills: Python, FastAPI, Redis, Celery")
    document.add_paragraph("Education: BS Computer Science")
    document.save(BytesIO())
    document.save(source_path)

    candidate_workspace = settings.workspace_dir / "candidates" / candidate_id
    candidate_workspace.mkdir(parents=True, exist_ok=True)
    (candidate_workspace / "resume_parsed.md").write_text(
        "legacy raw parser artifact",
        encoding="utf-8",
    )

    result = run_cv_extraction_flow(
        candidate_id=candidate_id,
        source_path=str(source_path),
        filename="resume.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert result["artifacts"]["structured_json_path"].endswith(
        "resume_structured.json"
    )
    assert result["artifacts"]["structured_markdown_path"].endswith("resume.md")
    structured_path = settings.workspace_dir / Path(
        result["artifacts"]["structured_json_path"]
    )
    markdown_path = settings.workspace_dir / Path(
        result["artifacts"]["structured_markdown_path"]
    )
    assert structured_path.exists()
    assert markdown_path.exists()
    assert "## Work Experience" in markdown_path.read_text(encoding="utf-8")
    assert not (markdown_path.parent / "resume_parsed.md").exists()


def test_cv_flow_stops_when_document_is_not_resume(tmp_path) -> None:
    settings.uploads_dir = tmp_path / "uploads"
    settings.workspace_dir = tmp_path / "workspace"
    upload_service.uploads_dir = settings.uploads_dir
    workspace_service.uploads_dir = settings.uploads_dir
    workspace_service.workspace_dir = settings.workspace_dir
    workspace_service.ensure_runtime_dirs()

    candidate_id = "candidate-not-resume"
    upload_dir = settings.uploads_dir / "resumes" / candidate_id
    upload_dir.mkdir(parents=True, exist_ok=True)
    source_path = upload_dir / "invoice.txt"
    source_path.write_text(
        "Invoice #1001\nAmount due: 1200\nPayment terms: net 30",
        encoding="utf-8",
    )

    with pytest.raises(InvalidResumeError, match="does not appear to be a resume"):
        run_cv_extraction_flow(
            candidate_id=candidate_id,
            source_path=str(source_path),
            filename="invoice.txt",
            content_type="text/plain",
        )

    candidate_workspace = settings.workspace_dir / "candidates" / candidate_id
    assert not (candidate_workspace / "resume_structured.json").exists()
    assert not (candidate_workspace / "resume.md").exists()
